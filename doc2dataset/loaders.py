"""
Document loaders for doc2dataset.

This module provides loaders for various document formats including
PDF, text, DOCX, and Markdown files.
"""

from __future__ import annotations

import os
import re
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, Iterator, List, Optional, Union


@dataclass
class Document:
    """
    Represents a loaded document.

    Attributes:
        content: The extracted text content.
        metadata: Metadata about the document (filename, pages, etc.).
        source: Original file path or URL.
        chunks: Optional list of document chunks.
    """

    content: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    source: str = ""
    chunks: List[str] = field(default_factory=list)

    @property
    def filename(self) -> str:
        """Get the filename from source."""
        return Path(self.source).name if self.source else ""

    @property
    def word_count(self) -> int:
        """Get approximate word count."""
        return len(self.content.split())

    def __len__(self) -> int:
        """Return content length."""
        return len(self.content)


class BaseLoader(ABC):
    """
    Abstract base class for document loaders.

    Subclasses implement loading logic for specific file formats.
    """

    @abstractmethod
    def load(self, path: Union[str, Path]) -> Document:
        """
        Load a document from a file path.

        Args:
            path: Path to the document file.

        Returns:
            Document object with extracted content.
        """
        pass

    @abstractmethod
    def supports(self, path: Union[str, Path]) -> bool:
        """
        Check if this loader supports the given file.

        Args:
            path: Path to check.

        Returns:
            True if this loader can handle the file.
        """
        pass

    def chunk(
        self,
        content: str,
        chunk_size: int = 2000,
        overlap: int = 200,
    ) -> List[str]:
        """
        Split content into chunks.

        Args:
            content: Text content to chunk.
            chunk_size: Target size of each chunk in characters.
            overlap: Overlap between consecutive chunks.

        Returns:
            List of text chunks.
        """
        if len(content) <= chunk_size:
            return [content]

        chunks = []
        start = 0

        while start < len(content):
            end = start + chunk_size

            if end < len(content):
                # Find a good break point
                for sep in ["\n\n", "\n", ". ", " "]:
                    break_point = content[start:end].rfind(sep)
                    if break_point > chunk_size // 2:
                        end = start + break_point + len(sep)
                        break

            chunk = content[start:end].strip()
            if chunk:
                chunks.append(chunk)

            start = end - overlap

        return chunks


class TextLoader(BaseLoader):
    """
    Loader for plain text files.

    Supports .txt files with various encodings.
    """

    EXTENSIONS = {".txt", ".text"}

    def __init__(self, encoding: str = "utf-8") -> None:
        """
        Initialize the text loader.

        Args:
            encoding: Text encoding to use.
        """
        self.encoding = encoding

    def supports(self, path: Union[str, Path]) -> bool:
        """Check if file is a text file."""
        return Path(path).suffix.lower() in self.EXTENSIONS

    def load(self, path: Union[str, Path]) -> Document:
        """Load a text file."""
        path = Path(path)

        try:
            content = path.read_text(encoding=self.encoding)
        except UnicodeDecodeError:
            # Try with different encoding
            content = path.read_text(encoding="latin-1")

        return Document(
            content=content,
            source=str(path),
            metadata={
                "format": "text",
                "encoding": self.encoding,
                "size_bytes": path.stat().st_size,
            },
        )


class MarkdownLoader(BaseLoader):
    """
    Loader for Markdown files.

    Extracts text content while preserving structure.
    """

    EXTENSIONS = {".md", ".markdown", ".mdown"}

    def __init__(self, strip_html: bool = True) -> None:
        """
        Initialize the Markdown loader.

        Args:
            strip_html: Whether to strip HTML tags.
        """
        self.strip_html = strip_html

    def supports(self, path: Union[str, Path]) -> bool:
        """Check if file is a Markdown file."""
        return Path(path).suffix.lower() in self.EXTENSIONS

    def load(self, path: Union[str, Path]) -> Document:
        """Load a Markdown file."""
        path = Path(path)
        content = path.read_text(encoding="utf-8")

        if self.strip_html:
            # Remove HTML tags
            content = re.sub(r"<[^>]+>", "", content)

        # Clean up multiple blank lines
        content = re.sub(r"\n{3,}", "\n\n", content)

        return Document(
            content=content.strip(),
            source=str(path),
            metadata={
                "format": "markdown",
                "size_bytes": path.stat().st_size,
            },
        )


class PDFLoader(BaseLoader):
    """
    Loader for PDF files.

    Uses PyMuPDF (fitz) for extraction.
    """

    EXTENSIONS = {".pdf"}

    def __init__(self, extract_images: bool = False) -> None:
        """
        Initialize the PDF loader.

        Args:
            extract_images: Whether to extract images (not implemented yet).
        """
        self.extract_images = extract_images

    def supports(self, path: Union[str, Path]) -> bool:
        """Check if file is a PDF."""
        return Path(path).suffix.lower() in self.EXTENSIONS

    def load(self, path: Union[str, Path]) -> Document:
        """Load a PDF file."""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise ImportError(
                "PyMuPDF not installed. Install with: pip install doc2dataset[pdf]"
            )

        path = Path(path)
        doc = fitz.open(str(path))

        pages = []
        for page_num, page in enumerate(doc):
            text = page.get_text()
            if text.strip():
                pages.append(text)

        content = "\n\n".join(pages)

        # Clean up common PDF artifacts
        content = self._clean_pdf_text(content)

        metadata = {
            "format": "pdf",
            "page_count": len(doc),
            "size_bytes": path.stat().st_size,
        }

        # Extract PDF metadata if available
        pdf_metadata = doc.metadata
        if pdf_metadata:
            for key in ["title", "author", "subject", "keywords"]:
                if pdf_metadata.get(key):
                    metadata[key] = pdf_metadata[key]

        doc.close()

        return Document(
            content=content,
            source=str(path),
            metadata=metadata,
        )

    def _clean_pdf_text(self, text: str) -> str:
        """Clean common PDF extraction artifacts."""
        # Fix hyphenated line breaks
        text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)
        # Normalize whitespace
        text = re.sub(r"[ \t]+", " ", text)
        # Clean up multiple newlines
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()


class DocxLoader(BaseLoader):
    """
    Loader for Microsoft Word documents.

    Uses python-docx for extraction.
    """

    EXTENSIONS = {".docx"}

    def supports(self, path: Union[str, Path]) -> bool:
        """Check if file is a DOCX."""
        return Path(path).suffix.lower() in self.EXTENSIONS

    def load(self, path: Union[str, Path]) -> Document:
        """Load a DOCX file."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError(
                "python-docx not installed. Install with: pip install doc2dataset[docx]"
            )

        path = Path(path)
        doc = DocxDocument(str(path))

        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)

        content = "\n\n".join(paragraphs)

        return Document(
            content=content,
            source=str(path),
            metadata={
                "format": "docx",
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "size_bytes": path.stat().st_size,
            },
        )


class JSONLoader(BaseLoader):
    """
    Loader for JSON files.

    Extracts text content from JSON structures.
    """

    EXTENSIONS = {".json"}

    def __init__(self, text_fields: Optional[List[str]] = None) -> None:
        """
        Initialize the JSON loader.

        Args:
            text_fields: List of field names to extract text from.
        """
        self.text_fields = text_fields or ["text", "content", "body", "description"]

    def supports(self, path: Union[str, Path]) -> bool:
        """Check if file is a JSON file."""
        return Path(path).suffix.lower() in self.EXTENSIONS

    def load(self, path: Union[str, Path]) -> Document:
        """Load a JSON file."""
        import json

        path = Path(path)
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)

        texts = []
        self._extract_texts(data, texts)

        content = "\n\n".join(texts)

        return Document(
            content=content,
            source=str(path),
            metadata={
                "format": "json",
                "size_bytes": path.stat().st_size,
            },
        )

    def _extract_texts(self, obj: Any, texts: List[str]) -> None:
        """Recursively extract text from JSON object."""
        if isinstance(obj, dict):
            for key, value in obj.items():
                if key.lower() in self.text_fields and isinstance(value, str):
                    if value.strip():
                        texts.append(value.strip())
                else:
                    self._extract_texts(value, texts)
        elif isinstance(obj, list):
            for item in obj:
                self._extract_texts(item, texts)


# Global loader registry
_LOADERS: List[BaseLoader] = [
    TextLoader(),
    MarkdownLoader(),
    PDFLoader(),
    DocxLoader(),
    JSONLoader(),
]


def load_document(
    path: Union[str, Path],
    loader: Optional[BaseLoader] = None,
) -> Document:
    """
    Load a document from a file path.

    Automatically selects the appropriate loader based on file extension.

    Args:
        path: Path to the document file.
        loader: Optional specific loader to use.

    Returns:
        Document object with extracted content.

    Raises:
        ValueError: If no suitable loader is found.

    Example:
        >>> doc = load_document("manual.pdf")
        >>> print(doc.content[:100])
    """
    path = Path(path)

    if loader:
        return loader.load(path)

    # Auto-select loader
    for ldr in _LOADERS:
        if ldr.supports(path):
            return ldr.load(path)

    raise ValueError(
        f"No loader found for file: {path}. "
        f"Supported extensions: {get_supported_extensions()}"
    )


def load_folder(
    folder: Union[str, Path],
    recursive: bool = True,
    extensions: Optional[List[str]] = None,
) -> Iterator[Document]:
    """
    Load all documents from a folder.

    Args:
        folder: Path to the folder.
        recursive: Whether to search subdirectories.
        extensions: List of extensions to include (e.g., [".pdf", ".txt"]).

    Yields:
        Document objects for each loaded file.

    Example:
        >>> for doc in load_folder("./documents"):
        ...     print(f"Loaded: {doc.filename}")
    """
    folder = Path(folder)

    if not folder.is_dir():
        raise ValueError(f"Not a directory: {folder}")

    # Determine which extensions to include
    if extensions:
        valid_extensions = set(ext.lower() for ext in extensions)
    else:
        valid_extensions = get_supported_extensions()

    # Get files
    if recursive:
        files = folder.rglob("*")
    else:
        files = folder.glob("*")

    for file_path in sorted(files):
        if not file_path.is_file():
            continue

        ext = file_path.suffix.lower()
        if ext not in valid_extensions:
            continue

        try:
            yield load_document(file_path)
        except Exception as e:
            # Log error but continue with other files
            print(f"Warning: Failed to load {file_path}: {e}")


def get_supported_extensions() -> set:
    """Get all supported file extensions."""
    extensions = set()
    for loader in _LOADERS:
        if hasattr(loader, "EXTENSIONS"):
            extensions.update(loader.EXTENSIONS)
    return extensions


def register_loader(loader: BaseLoader) -> None:
    """
    Register a custom loader.

    Args:
        loader: Loader instance to register.
    """
    _LOADERS.insert(0, loader)  # Insert at beginning for priority
